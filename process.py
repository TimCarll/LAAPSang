from xlrd import open_workbook
import xlrd
from docx import Document
from docx.shared import Inches
import unidecode 
import glob, os
import re

execfile('global.py')

def validate(number):
	if isinstance(number, (int, float)):
		return number
	elif number is None:
		return -1
	elif number == "":
		return -1
	else:
		match = re.match('^([0-9.]+)',number)
		if match:
			return float(match.group(1))
		
		match = re.match('^<([0-9.]+)', number)
		if match:
			return float(match.group(1))

		match = re.match('^>([0-9.]+)', number)
		if match:
			return float(match.group(1))


# this file has a function that processes one file, and returns a hash with all the values
def readFile(filename):
	book = open_workbook(filename,on_demand=True)
	for name in book.sheet_names():
    		if name.endswith('Lab'):
			sheet = book.sheet_by_name(name)
	d = dict()
	s = []
	drugs = [sheet.cell_value(3, 0),sheet.cell_value(3,1),sheet.cell_value(3,2)]
	for i in drugs:
		if i != "":
			s.append(str(i))	
	d["DRUG"] = "+".join(s)
	d["PT_INIT"] = sheet.cell_value(3, 5)
	d["PT_MRN"] = sheet.cell_value(4, 5)
	d["DATE"] = sheet.cell_value(5, 5)
	d["LAPTT_R"] = validate(sheet.cell_value(11, 1))
	d["PTTMX_R"] = sheet.cell_value(14, 1)
	d["LTT_R"] = sheet.cell_value(12, 1)
	d["LTTHEP_R"] = sheet.cell_value(13, 1)
	d["LTTHEP_P"] = sheet.cell_value(13, 4)
	d["LAPTT_U"] = sheet.cell_value(11, 3)
	d["LTT_L"] = sheet.cell_value(12, 2)
	d["LTT_U"] = sheet.cell_value(12, 3)
	d["LTTHEP_U"] = sheet.cell_value(13, 3)
	d["PTTMX_U"] = sheet.cell_value(14, 3)
	d["LAPTT_P"] = sheet.cell_value(11, 4)
	d["LTT_P"] = sheet.cell_value(12, 4)
	d["LTTHEP_P"] = sheet.cell_value(13, 4)
	d["PTTMX_P"] = sheet.cell_value(14, 4)

	d["PNP_R"] = validate(sheet.cell_value(20, 1))
	d["PNP_R_LYS"] = validate(sheet.cell_value(20, 2))
	d["PNP_SD"] = validate(sheet.cell_value(20, 4))
	d["PNP_U"] = validate(sheet.cell_value(20, 6))

	# have to fix this in excel sheet, or else PNP_SD defaults to 42.0
	if d["PNP_R"] == "":
		d["PNP_SD"] = 0
		d["PNP_R"]  = 0
	
	d["DRVVS_R"] = validate(sheet.cell_value(25, 1))
	d["DRVVMX_R"] = sheet.cell_value(26, 1)
	d["DRVVC_R"] = sheet.cell_value(27, 1)
	d["PCTCO_R"] = sheet.cell_value(28, 1)
	d["DRVVS_U"] = sheet.cell_value(25, 3)
	d["DRVVMX_U"] = sheet.cell_value(26, 3)
	d["PCTCO_U"] = sheet.cell_value(28, 3)
	d["DRVVS_P"] = sheet.cell_value(25, 4)
	d["DRVVMX_P"] = sheet.cell_value(26, 4)
	d["PCTCO_P"] = sheet.cell_value(28, 4)
	d["DPTS_R"] = validate(sheet.cell_value(33, 1))
	d["DPTMX_R"] = sheet.cell_value(34, 1)
	d["DPTC_R"] = sheet.cell_value(35, 1)
	d["DPTCOR_R"] = sheet.cell_value(36, 1)
	d["DPTS_U"] = sheet.cell_value(33, 3)
	d["DPTMX_U"] = sheet.cell_value(34, 3)
	d["DPTCOR_U"] = sheet.cell_value(36, 3)
	d["DPTS_P"] = sheet.cell_value(33, 4)
	d["DPTMX_P"] = sheet.cell_value(34, 4)
	d["DPTCOR_P"] = sheet.cell_value(36, 4)


	# Staclot-LA added 6/29/2017
	if validate(sheet.cell_value(11,6))!=-1:
		d["STACLOT-LA1"] = validate(sheet.cell_value(11,6))
		d["STACLOT-LA1-ULN"] = validate(sheet.cell_value(11,7))
		d["STACLOT-LA1-95"] = validate(sheet.cell_value(11,8))
		d["STACLOT-LA2"] = validate(sheet.cell_value(12,6))
		d["STACLOT-LA2-ULN"] = validate(sheet.cell_value(12,7))
		d["STACLOT-LA2-95"] = validate(sheet.cell_value(12,8))
		d["STACLOT-CORR"] = validate(sheet.cell_value(13,6))	
		d["STACLOT-CORR-ULN"] = validate(sheet.cell_value(13,7))
		d["STACLOT-CORR-95"] = validate(sheet.cell_value(13,8))
		d["STACLOT-USED"] = True
	else:
		d["STACLOT-LA1"] = -1
		d["STACLOT-LA1-ULN"] = -1
		d["STACLOT-LA1-95"] = -1
		d["STACLOT-LA2"] = -1
		d["STACLOT-LA2-ULN"] = -1
		d["STACLOT-LA2-95"] = -1
		d["STACLOT-CORR"] = -1
		d["STACLOT-CORR-ULN"] = -1
		d["STACLOT-CORR-95"] = -1
		d["STACLOT-USED"] = False


	try:
		d["AG_1_R"] = validate(sheet.cell_value(41,1)) 
		d["AG_2_R"] = validate(sheet.cell_value(42,1)) 
		d["AG_3_R"] = validate(sheet.cell_value(43,1)) 
		d["AG_4_R"] = validate(sheet.cell_value(44,1)) 
		d["AG_5_R"] = validate(sheet.cell_value(48,1)) 
		d["AG_6_R"] = validate(sheet.cell_value(49,1)) 
		d["AG_7_R"] = validate(sheet.cell_value(50,1)) 
		d["AG_8_R"] = validate(sheet.cell_value(51,1)) 

		d["AG_1_U"] = validate(sheet.cell_value(41,3)) 
		d["AG_2_U"] = validate(sheet.cell_value(42,3)) 
		d["AG_3_U"] = validate(sheet.cell_value(43,3)) 
		d["AG_4_U"] = validate(sheet.cell_value(44,3)) 
		d["AG_5_U"] = validate(sheet.cell_value(48,3)) 
		d["AG_6_U"] = validate(sheet.cell_value(49,3)) 
		d["AG_7_U"] = validate(sheet.cell_value(50,3)) 
		d["AG_8_U"] = validate(sheet.cell_value(51,3)) 
	except:
		d["AG_1_R"] = -1
		d["AG_2_R"] = -1
		d["AG_3_R"] = -1
		d["AG_4_R"] = -1
		d["AG_5_R"] = -1
		d["AG_6_R"] = -1
		d["AG_7_R"] = -1
		d["AG_8_R"] = -1

		d["AG_1_U"] = -1
		d["AG_2_U"] = -1
		d["AG_3_U"] = -1
		d["AG_4_U"] = -1
		d["AG_5_U"] = -1
		d["AG_6_U"] = -1
		d["AG_7_U"] = -1
		d["AG_8_U"] = -1
	return d

def writeFile(txt, filename):
	document = Document()
	document.add_paragraph(txt)	
	document.save(filename)


def extractConclusion(filename):
	if filename is None:
		return ""
	f = os.path.join(REPORTDIR,filename)
	document = Document(f)
	s = []
	flag = False
	for p in document.paragraphs:
		cleanedText = unidecode.unidecode(p.text)
		cleanedText = cleanedText.replace("\n"," ")
		cleanedText = cleanedText.replace("\t"," ")
		if (cleanedText.startswith("Conclusion") or cleanedText.startswith("Note")) and (not (cleanedText.startswith("Jonathan") or (cleanedText.startswith("Angela")))):
			s.append(cleanedText)
	return "\n".join(s)

def findFiles(filename):
	strippedFilename = os.path.basename(filename)
	a = strippedFilename.split(" ");
	s = []
	for file in os.listdir(REPORTDIR):
		if (file.endswith(".docx") or (file.endswith(".doc"))) and file.startswith(a[0]+" "+a[1]):
			s.append(file)
	s.sort()
	return s

def findLastFile(filename):
	s = findFiles(filename)
	if len(s)>0:
		return s[len(s)-1]
	else:
		return None


def dumpDataHeader(filehandle):
    header = "\"DRUG\", \"PT_INIT\", \"PT_MRN\", \"FILE\", \"LAPTT_R\", \"PTTMX_R\", \"LTT_R\", \"LTTHEP_R\", \"LTTHEP_P\", \"LAPTT_U\", \"LTT_L\", \"LTT_U\", \"LTTHEP_U\", \"PTTMX_U\", \"LAPTT_P\", \"LTT_P\", \"LTTHEP_P\", \"PTTMX_P\", \"PNP_R\", \"PNP_R_LYS\", \"PNP_SD\", \"PNP_U\", \"DRVVS_R\", \"DRVVMX_R\", \"DRVVC_R\", \"PCTCO_R\", \"DRVVS_U\", \"DRVVMX_U\", \"PCTCO_U\", \"DRVVS_P\", \"DRVVMX_P\", \"PCTCO_P\", \"DPTS_R\", \"DPTMX_R\", \"DPTC_R\", \"DPTCOR_R\", \"DPTS_U\", \"DPTMX_U\", \"DPTCOR_U\", \"DPTS_P\", \"DPTMX_P\", \"DPTCOR_P\", \"AG_1_R\", \"AG_2_R\", \"AG_3_R\", \"AG_4_R\", \"AG_5_R\", \"AG_6_R\", \"AG_7_R\", \"AG_8_R\", \"AG_1_U\", \"AG_2_U\", \"AG_3_U\", \"AG_4_U\", \"AG_5_U\", \"AG_6_U\", \"AG_7_U\", \"AG_8_U\",\"STACLOT-USED\",\"STACLOT-LA1\",\"STACLOT-LA1-ULN\",\"STACLOT-LA1-95\",\"STACLOT-LA2\",\"STACLOT-LA2-ULN\",\"STACLOT-LA2-95\",\"STACLOT-CORR\",\"STACLOT-CORR-ULN\",\"STACLOT-CORR-95\""
    header = "\"Drug\", \"Patient Initials\", \"Patient MRN\", \"File\",  \"LAPTT Abnormal\", \"DRVVT Abnormal\", \"DPT Abnormal\", \"Antigenic Testing 1 Abnormal\", \"Antigenic Testing 2 Abnormal\", \"LAPTT Screening Result\", \"PTT Mixing Result\", \"LTT Result\", \"LTTHEP Result\", \"LTTHEP Prolongation\", \"LAPTT Upper Limit of Normal\", \"LTT Lower Limit of Normal\", \"LTT Upper Limit of Normal\", \"LTTHEP Upper Limit of Normal\", \"PTTMX Upper Limit of Normal\", \"LAPTT Prolongation\", \"LTT Prolongation\", \"LTTHEP Prolongation\", \"PTTMX Prolongation\", \"PNP Result\", \"PNP Result using Lysate\", \"PNP_SD\", \"PNP Upper Limit of Normal\", \"DRVVT Screening Result\", \"DRVVT Mixing Result\", \"DRVVT Result\", \"DRVVT Percent Correction Result\", \"DRVVT Upper Limit of Normal\", \"DRVVT Mixing Upper Limit of Normal\", \"DRVVT % Correction Upper Limit of Normal\", \"DRVVT Prolongation\", \"DRVVT Mixing Prolongation\", \"DRVVT % Correction Prolongation\", \"DPT Screening Result\", \"DPT Mixing Result\", \"DPT Correction Result\", \"DPT % Correction Result\", \"DPT Screen Upper Limit of Normal\", \"DPT Mixing Upper Limit of Normal\", \"DPT % Correction Upper Limit of Normal\", \"DPT Screening Prolongation\", \"DPT Mixing Prolongation\", \"DPT % Correction Prolongation\", \"aCA IgG Result\", \"aCA IgM Result\", \"beta2 IgG Result\", \"beta2 IgM Result\", \"aCA IgA Result\", \"beta2 IgA Result\", \"APS/PT IgG Result\", \"APS/PT IgM Result\", \"AG_1 Upper Limit of Normal\", \"aCA IgM Upper Limit of Normal\", \"beta2 IgG Upper Limit of Normal\", \"beta2 IgM Upper Limit of Normal\", \"aCA IgA Upper Limit of Normal\", \"beta2 IgA Upper Limit of Normal\", \"APS/PT IgG Upper Limit of Normal\", \"APS/PT IgM Upper Limit of Normal\",\"Staclot used\", \"Staclot LA1 Pt Result\",\"Staclot LA1 ULN\",\"Staclot LA1 95th\",\"Staclot LA2 Pt Result\",\"Staclot LA2 ULN\",\"Staclot LA2 95th\",\"Staclot Corr Pt Result\",\"Staclot Corr ULN\",\"Staclot Corr 95th\""
    filehandle.write(header + "\n")
    return

def dumpData(file, filehandle):
    d = readFile(os.path.join(REPORTDIR,file))
 
    if d["AG_1_R"]>d["AG_1_U"] or  d["AG_2_R"]>d["AG_2_U"] or d["AG_3_R"]>d["AG_3_U"] or  d["AG_4_R"]>d["AG_4_U"]:  
        agtesting1 = True
    else:
        agtesting1 = False
    if d["AG_5_R"]>d["AG_5_U"] or  d["AG_6_R"]>d["AG_6_U"] or d["AG_7_R"]>d["AG_7_U"] or  d["AG_8_R"]>d["AG_8_U"]:  
        agtesting2 = True
    else:
        agtesting2 = False

    a = [d["DRUG"], d["PT_INIT"], d["PT_MRN"], "\""+file+"\"", d["PNP_SD"]>d["PNP_U"] , d["PCTCO_R"]>d["PCTCO_U"], d["DPTCOR_R"]>d["DPTCOR_U"], agtesting1, agtesting2, d["LAPTT_R"], d["PTTMX_R"], d["LTT_R"], d["LTTHEP_R"], d["LTTHEP_P"], d["LAPTT_U"], d["LTT_L"], d["LTT_U"], d["LTTHEP_U"], d["PTTMX_U"], d["LAPTT_P"], d["LTT_P"], d["LTTHEP_P"], d["PTTMX_P"], d["PNP_R"], d["PNP_R_LYS"], d["PNP_SD"], d["PNP_U"], d["DRVVS_R"], d["DRVVMX_R"], d["DRVVC_R"], d["PCTCO_R"], d["DRVVS_U"], d["DRVVMX_U"], d["PCTCO_U"], d["DRVVS_P"], d["DRVVMX_P"], d["PCTCO_P"], d["DPTS_R"], d["DPTMX_R"], d["DPTC_R"], d["DPTCOR_R"], d["DPTS_U"], d["DPTMX_U"], d["DPTCOR_U"], d["DPTS_P"], d["DPTMX_P"], d["DPTCOR_P"], d["AG_1_R"], d["AG_2_R"], d["AG_3_R"], d["AG_4_R"], d["AG_5_R"], d["AG_6_R"], d["AG_7_R"], d["AG_8_R"], d["AG_1_U"], d["AG_2_U"], d["AG_3_U"], d["AG_4_U"], d["AG_5_U"], d["AG_6_U"], d["AG_7_U"], d["AG_8_U"],d["STACLOT-USED"], d["STACLOT-LA1"], d["STACLOT-LA1-ULN"], d["STACLOT-LA1-95"], d["STACLOT-LA2"], d["STACLOT-LA2-ULN"], d["STACLOT-LA2-95"], d["STACLOT-CORR"], d["STACLOT-CORR-ULN"], d["STACLOT-CORR-95"]]

    b = map(lambda x: str(x), a)

    filehandle.write(",".join(b) + "\n")
   
    return
    
def dumpDataReportDir(outputfilename):
    targetfile = open(outputfilename,"w")
    dumpDataHeader(targetfile)
    for file in os.listdir(REPORTDIR):
        if (file.endswith(".xlsx") or (file.endswith(".xls"))):
            print "Processing file: " + file
#            try:
            dumpData(file, targetfile)    
#            except:
#                print "Error processing file: " + file
    targetfile.close()
    return



